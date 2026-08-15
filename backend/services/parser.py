import fitz  # PyMuPDF
import docx
from io import BytesIO
from fastapi import UploadFile


async def parse_txt(file: UploadFile) -> str:
    content = await file.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 decodes any byte, so we rarely reach here, but guarantee a return.
    return content.decode("utf-8", errors="replace")


async def parse_pdf(file: UploadFile) -> str:
    content = await file.read()
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        raise ValueError(f"Error parsing PDF: {e}")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


async def parse_docx(file: UploadFile) -> str:
    content = await file.read()
    try:
        doc = docx.Document(BytesIO(content))
        parts = [para.text for para in doc.paragraphs]
        # python-docx excludes table cells from .paragraphs, so walk tables too —
        # resumes, forms, and invoices often keep their content in tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Error parsing DOCX: {e}")


async def parse_file(file: UploadFile) -> str:
    if not file.filename or "." not in file.filename:
        raise ValueError("File has no recognizable extension")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return await parse_pdf(file)
    elif ext == "docx":
        return await parse_docx(file)
    elif ext == "txt":
        return await parse_txt(file)
    else:
        raise ValueError(f"Unsupported file type: {ext}")